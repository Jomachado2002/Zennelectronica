#!/usr/bin/env python3
"""Genera el contrato formal de licencia e implementación del software."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/josiasnicolas02gmail.com/Documents/Zennelectronica/CONTRATO_LICENCIA_ECOMMERCE_ERP_ZENN.docx"

PRIMARY = RGBColor(0x0B, 0x3D, 0x91)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x44, 0x44, 0x44)
HEADER_BG = "0B3D91"
ZEBRA = "F2F5FA"
FONT = "Times New Roman"


def set_run(run, size=11, bold=False, color=DARK, italic=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def p(doc, text="", size=11, bold=False, align="justify", space_after=8, space_before=0, italic=False, color=DARK):
    para = doc.add_paragraph()
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "left":
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if text:
        run = para.add_run(text)
        set_run(run, size=size, bold=bold, color=color, italic=italic)
    return para


def add_mixed(para, parts):
    """parts: list of (text, bold, italic optional)"""
    for part in parts:
        text = part[0]
        bold = part[1] if len(part) > 1 else False
        italic = part[2] if len(part) > 2 else False
        size = part[3] if len(part) > 3 else 11
        run = para.add_run(text)
        set_run(run, size=size, bold=bold, italic=italic)


def heading(doc, text, level=1):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(text)
    set_run(run, size=13 if level == 1 else 11.5, bold=True, color=PRIMARY)
    if level == 1:
        pPr = para._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "B0C4DE")
        pbdr.append(bottom)
        pPr.append(pbdr)
    return para


def bullet(doc, text, bold_prefix=None):
    para = doc.add_paragraph(style="List Bullet")
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = para.add_run(bold_prefix)
        set_run(r1, bold=True)
        r2 = para.add_run(text)
        set_run(r2)
    else:
        r = para.add_run(text)
        set_run(r)
    return para


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], HEADER_BG)
        para = hdr[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        set_run(run, size=10, bold=True, color=RGBColor(255, 255, 255))
    for r_i, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if r_i % 2 == 1:
                set_cell_bg(cells[i], ZEBRA)
            para = cells[i].paragraphs[0]
            run = para.add_run(val)
            set_run(run, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)

    # PORTADA / TÍTULO
    p(doc, "CONTRATO DE LICENCIA DE USO, IMPLEMENTACIÓN", size=14, bold=True, align="center", space_after=2, color=PRIMARY)
    p(doc, "Y PRESTACIÓN DE SERVICIOS DE SOFTWARE", size=14, bold=True, align="center", space_after=6, color=PRIMARY)
    p(
        doc,
        "“Plataforma E-commerce + ERP (base Zenn)”",
        size=12,
        bold=True,
        align="center",
        space_after=14,
        italic=True,
    )

    p(
        doc,
        "En la ciudad de Asunción, República del Paraguay, a los veintidós (22) días del mes de julio del año dos mil veintiséis (2026), se celebra el presente Contrato de Licencia de Uso, Implementación y Prestación de Servicios de Software (en adelante, el “Contrato”), entre las partes que a continuación se individualizan:",
        space_after=12,
    )

    heading(doc, "CLÁUSULA PRIMERA – PARTES")

    para = p(doc, "", space_after=6)
    add_mixed(
        para,
        [
            ("1.1. EL PROVEEDOR: ", True),
            (
                "ZENN E.A.S. UNIPERSONAL, con RUC N.º 80162448-7, con domicilio en Teodoro S. Mongelós casi Radio Operadores del Chaco N.º 3934, Asunción, Paraguay, representada en este acto por su representante legal, el Sr. JOSÍAS NICOLÁS MACHADO ROLÓN, con Cédula de Identidad Civil N.º 5799267, en adelante “ZENN” o el “PROVEEDOR”.",
                False,
            ),
        ],
    )

    para = p(doc, "", space_after=6)
    add_mixed(
        para,
        [
            ("1.2. EL CLIENTE: ", True),
            (
                "[RAZÓN SOCIAL], con RUC N.º [COMPLETAR], con domicilio en [COMPLETAR DOMICILIO], representada en este acto por [NOMBRE DEL REPRESENTANTE LEGAL], con Cédula de Identidad Civil N.º [COMPLETAR], en adelante el “CLIENTE”.",
                False,
            ),
        ],
    )

    p(
        doc,
        "ZENN y el CLIENTE podrán ser denominados individualmente como “Parte” y conjuntamente como “las Partes”.",
        space_after=10,
    )

    heading(doc, "CLÁUSULA SEGUNDA – OBJETO")
    p(
        doc,
        "2.1. El objeto del presente Contrato es la otorgación por parte de ZENN al CLIENTE de una licencia de uso del software denominado “Plataforma E-commerce + ERP (base Zenn)” (en adelante, el “Software”), la implementación personalizada del mismo para la empresa del CLIENTE, la entrega del código fuente correspondiente, la capacitación incluida, una bolsa de horas de desarrollo adicional, y la regulación de los servicios posteriores de mantenimiento, soporte y mejoras, conforme a los términos aquí establecidos.",
    )
    p(
        doc,
        "2.2. El Software constituye una plataforma de comercio electrónico integrada con sistema de gestión comercial (ERP) y herramientas de automatización, cuya funcionalidad completa se detalla en el Anexo A (Inventario de Funcionalidades), el cual forma parte integrante del presente Contrato.",
    )

    heading(doc, "CLÁUSULA TERCERA – LICENCIA DE USO Y PROPIEDAD INTELECTUAL")
    p(
        doc,
        "3.1. ZENN declara ser el autor y titular de los derechos patrimoniales sobre el Software y su base tecnológica, creados bajo la dirección y autoría del CEO y representante legal de ZENN, Sr. Josías Nicolás Machado Rolón, sin perjuicio de componentes de terceros de código abierto o servicios externos (Bancard, MongoDB, Firebase, etc.).",
    )
    p(
        doc,
        "3.2. Por el presente Contrato, ZENN otorga al CLIENTE una licencia de uso perpetua, onerosa, no exclusiva e intransferible (salvo lo previsto en el punto 3.5), limitada al uso propio de la empresa del CLIENTE, para operar su negocio a través del Software.",
    )
    p(
        doc,
        "3.3. El CLIENTE tendrá derecho a: (i) utilizar el Software en su operación comercial; (ii) acceder y conservar el código fuente en un repositorio GitHub a su nombre; (iii) modificar el código por sí o mediante terceros, bajo su exclusiva responsabilidad.",
    )
    p(
        doc,
        "3.4. Queda expresamente prohibido al CLIENTE: (i) vender, ceder, sublicenciar, alquilar o comercializar el Software o su código fuente como producto o servicio a terceros; (ii) explotar el Software como solución llave en mano para otras empresas; (iii) utilizar la licencia fuera del ámbito del uso propio de su empresa. El incumplimiento de esta cláusula constituirá causal grave de resolución, sin perjuicio de las acciones legales por daños y perjuicios y violación de propiedad intelectual.",
    )
    p(
        doc,
        "3.5. En caso de venta, fusión o cesión del fondo de comercio o de la empresa del CLIENTE, la licencia podrá transferirse al adquirente únicamente para continuar la operación del mismo negocio, debiendo notificarse a ZENN por escrito. Ello no genera obligación de pago adicional a ZENN por la sola transferencia, sin perjuicio de los servicios de soporte que el adquirente desee contratar.",
    )
    p(
        doc,
        "3.6. ZENN conserva el derecho de desarrollar, licenciar y comercializar sistemas iguales, similares o derivados de la misma base tecnológica a otros clientes, por tratarse de una licencia no exclusiva de uso propio.",
    )
    p(
        doc,
        "3.7. La presente operación no implica cesión total de derechos de autor ni exclusividad a favor del CLIENTE, sino licencia de uso con entrega de código fuente en los términos de esta cláusula.",
    )

    heading(doc, "CLÁUSULA CUARTA – PRECIO Y FORMA DE PAGO")
    p(
        doc,
        "4.1. El precio total de la licencia, implementación inicial (Fase 1 y Fase 2), diseño, configuración de servidores/entornos, entrega de código fuente, diez (10) horas de capacitación y treinta (30) horas de desarrollo adicional incluidas es de USD 5.000 (cinco mil dólares estadounidenses), Impuesto al Valor Agregado (IVA) incluido.",
    )
    p(doc, "4.2. El precio se abonará exclusivamente en dólares estadounidenses (USD), conforme al siguiente cronograma:")

    add_table(
        doc,
        ["Hito", "Porcentaje", "Monto (USD)", "Momento de pago"],
        [
            ["Anticipo / inicio", "30%", "1.500", "Al firmar el Contrato, para iniciar la implementación"],
            ["Cierre Fase 1", "30%", "1.500", "Contra Acta de Aceptación de Fase 1"],
            ["Cierre Fase 2", "40%", "2.000", "Contra Acta de Aceptación de Fase 2 (entrega final)"],
            ["TOTAL", "100%", "5.000", "IVA incluido"],
        ],
    )

    p(
        doc,
        "4.3. Cada pago deberá realizarse dentro de los cinco (5) días hábiles siguientes a la emisión de la factura correspondiente y/o a la firma del Acta de Aceptación del hito, según corresponda. Los datos bancarios de ZENN serán informados al CLIENTE por escrito.",
    )
    p(
        doc,
        "4.4. El precio indicado no incluye costos de infraestructura ni servicios de terceros (dominio, correo SMTP, VPS, Bancard, excedentes de MongoDB Atlas o Firebase, publicidad, etc.), los cuales serán a exclusivo cargo del CLIENTE.",
    )

    heading(doc, "CLÁUSULA QUINTA – FASES DE IMPLEMENTACIÓN")
    p(
        doc,
        "5.1. La implementación se ejecutará en dos (2) fases. El detalle funcional completo del Software se encuentra en el Anexo A. Todo requerimiento que no forme parte del Anexo A se considerará fuera de alcance y se gestionará mediante la bolsa de treinta (30) horas o mediante cotización / tarifa horaria, conforme a este Contrato.",
    )

    para = p(doc, "", space_after=4)
    add_mixed(para, [("5.2. FASE 1 – Primer borrador / visualización (plazo: un (1) mes):", True)])
    p(
        doc,
        "Tiene por objeto que el CLIENTE pueda visualizar el sistema con su identidad de marca y el catálogo de productos a vender. Incluye, enunciativamente:",
        space_after=4,
    )
    bullet(doc, "Creación y configuración de base de datos y entornos necesarios.")
    bullet(doc, "Montaje y configuración de servidores / hosting de la web y componentes asociados por parte de ZENN.")
    bullet(doc, "Diseño e implementación visual con logo, paleta de colores y tipografías provistos por el CLIENTE.")
    bullet(doc, "Publicación de un sitio navegable que permita ver productos, categorías, estructura comercial y tipología del sistema (primer borrador).")
    bullet(doc, "Uso eventual de horas de la bolsa de treinta (30) horas, si resultare necesario para mostrar al CLIENTE alguna diferenciación puntual.")
    p(
        doc,
        "En Fase 1 no se exige necesariamente la habilitación productiva de cobros Bancard, la cual corresponde a la Fase 2.",
        space_after=8,
    )

    para = p(doc, "", space_after=4)
    add_mixed(para, [("5.3. FASE 2 – Entrega final / puesta en marcha completa:", True)])
    p(
        doc,
        "Tiene por objeto dejar el Software operativo conforme al inventario del Anexo A, incluyendo de manera enunciativa y no limitativa:",
        space_after=4,
    )
    bullet(doc, "Integración y configuración de Bancard y flujos de cobro online.")
    bullet(doc, "Back office / ERP, automatizaciones, herramientas comerciales y demás funcionalidades del Anexo A.")
    bullet(doc, "Ajustes específicos del CLIENTE mediante la bolsa de hasta treinta (30) horas de desarrollo adicional.")
    bullet(doc, "Entrega del código fuente en repositorio GitHub y transferencia de accesos.")
    bullet(doc, "Capacitación incluida de hasta diez (10) horas.")
    p(
        doc,
        "El plazo de la Fase 2 dependerá de las necesidades de implementación y mejoras que el CLIENTE requiera y autorice por escrito. La Fase 2 constituye la entrega final del proyecto objeto de los USD 5.000.",
        space_after=8,
    )

    p(
        doc,
        "5.4. ZENN montará y configurará el proyecto durante la vigencia de la implementación. Al cierre, transferirá al CLIENTE los accesos, credenciales y documentación razonable de operación. La titularidad/contratación de dominio, correo, VPS, Bancard, MongoDB y Firebase será del CLIENTE.",
    )

    heading(doc, "CLÁUSULA SEXTA – ACTAS DE ACEPTACIÓN")
    p(
        doc,
        "6.1. Al concluir cada Fase, ZENN notificará al CLIENTE la disponibilidad para revisión. El CLIENTE contará con quince (15) días corridos para suscribir el Acta de Aceptación o formular observaciones por escrito vinculadas al alcance contractual.",
    )
    p(
        doc,
        "6.2. Si el CLIENTE no responde dentro de ese plazo, la Fase se considerará tácitamente aceptada, pudiendo ZENN emitir la factura del hito correspondiente.",
    )
    p(
        doc,
        "6.3. Las observaciones válidas serán únicamente aquellas que se refieran a funcionalidades incluidas en el Anexo A o en requerimientos escritos aprobados, que no funcionen conforme a lo acordado (bugs/errores). Los pedidos de nuevas funciones no constituyen observación de aceptación, sino mejora o cambio de alcance.",
    )

    heading(doc, "CLÁUSULA SÉPTIMA – BOLSA DE TREINTA (30) HORAS DE DESARROLLO ADICIONAL")
    p(
        doc,
        "7.1. Dentro del precio de USD 5.000 se incluyen treinta (30) horas de desarrollo adicional destinadas a features, diferenciaciones o implementaciones que el sistema actualmente no contemple en el Anexo A, o cambios diferenciales requeridos por el CLIENTE.",
    )
    p(
        doc,
        "7.2. Dichas horas podrán utilizarse en Fase 1 y/o Fase 2, según necesidad de lo que deba mostrarse o entregarse al CLIENTE. Las horas no consumidas al cierre de Fase 2 podrán utilizarse durante los cuarenta (40) días posteriores a la aceptación de Fase 2.",
    )
    p(
        doc,
        "7.3. El consumo de horas se gestionará mediante planillas de avance y confirmación previa del autorizador que el CLIENTE designe por escrito. Sin autorización previa, ZENN no estará obligado a ejecutar trabajos con cargo a la bolsa de horas.",
    )
    p(
        doc,
        "7.4. Agotada la bolsa de treinta (30) horas, o vencido el plazo de uso de las no consumidas, todo desarrollo adicional se cotizará y facturará a la tarifa horaria de la Cláusula Novena.",
    )
    p(
        doc,
        "7.5. Las treinta (30) horas no cubren corrección de bugs propios de lo ya contemplado en el Anexo A durante el período de garantía de la Cláusula Octava, los cuales se corrigen sin cargo en ese período.",
    )

    heading(doc, "CLÁUSULA OCTAVA – GARANTÍA POST-ENTREGA (40 DÍAS)")
    p(
        doc,
        "8.1. A partir de la firma del Acta de Aceptación de Fase 2 (o su aceptación tácita), el CLIENTE gozará de cuarenta (40) días corridos de garantía para la corrección de bugs o errores de la plataforma relativos a funcionalidades incluidas en el Anexo A o en cambios autorizados ya implementados.",
    )
    p(
        doc,
        "8.2. Esta garantía no consume la bolsa de treinta (30) horas ni genera cargo por tarifa horaria, siempre que se trate de fallas (lo acordado no funciona), y no de nuevas funcionalidades o cambios de alcance.",
    )
    p(
        doc,
        "8.3. Quedan excluidos de la garantía: fallas de infraestructura de terceros, mala operación del CLIENTE, modificaciones hechas por terceros sin intervención de ZENN, cambios de APIs externas (Bancard, etc.), y fuerza mayor.",
    )

    heading(doc, "CLÁUSULA NOVENA – MANTENIMIENTO, SOPORTE Y MEJORAS (POST-PROYECTO)")
    p(
        doc,
        "9.1. Finalizada la implementación y/o agotadas las prestaciones incluidas en los USD 5.000, los servicios de mantenimiento, soporte, mejoras, nuevas implementaciones y cualquier trabajo no contemplado en el código/alcance del Anexo A se prestarán por demanda, a la tarifa de USD 15 (quince dólares estadounidenses) por hora.",
    )
    p(
        doc,
        "9.2. La unidad mínima de facturación por intervención será de una (1) hora. Los trabajos se autorizarán previamente por el autorizador designado por el CLIENTE, se registrarán en planillas y se facturarán mensualmente conforme a las horas efectivamente adeudadas en el período.",
    )
    p(
        doc,
        "9.3. No se incluyen en esta tarifa horaria (por estar cubiertos en el precio de USD 5.000): el diseño e implementación inicial de Fase 1 y Fase 2, la licencia de uso, la entrega de código fuente, la configuración inicial de servidores/entornos, las diez (10) horas de capacitación incluidas y las treinta (30) horas de desarrollo adicional.",
    )
    p(
        doc,
        "9.4. Si el CLIENTE requiere capacitación adicional a las diez (10) horas incluidas, dicha capacitación se cobrará a USD 15 por hora.",
    )
    p(
        doc,
        "9.5. El CLIENTE podrá, si lo desea, contratar a otro desarrollador para modificar el código. Ello no genera responsabilidad de ZENN por trabajos de terceros. Mientras exista vínculo de soporte con ZENN, se recomienda canalizar los cambios a través de ZENN para preservar la integridad del sistema.",
    )

    heading(doc, "CLÁUSULA DÉCIMA – CAPACITACIÓN")
    p(
        doc,
        "10.1. ZENN incluirá hasta diez (10) horas de capacitación al equipo del CLIENTE, dentro del precio de USD 5.000, a coordinarse a partir de la Fase 2 o según acuerdo de las Partes.",
    )
    p(
        doc,
        "10.2. Las horas de capacitación adicionales, incluyendo las que se requieran en un eventual traspaso por terminación del vínculo de soporte, serán onerosas a USD 15 por hora.",
    )

    heading(doc, "CLÁUSULA DÉCIMA PRIMERA – INFRAESTRUCTURA Y LIMITACIÓN DE GARANTÍAS OPERATIVAS")
    p(
        doc,
        "11.1. Estarán a cargo del CLIENTE: dominio, correo corporativo/SMTP para envío de emails, servidor VPS u otros servicios necesarios para procesos pesados (PDF, workers, etc.), contrato y credenciales Bancard, y cuentas de MongoDB Atlas, Firebase u otros proveedores cloud, incluyendo eventuales costos por superar planes freemium.",
    )
    p(
        doc,
        "11.2. ZENN no garantiza un volumen mínimo de ventas, resultados comerciales, posicionamiento SEO ni retorno de inversión. El Software es una herramienta; el resultado de negocio depende del CLIENTE (precios, stock, marketing, atención, logística, etc.).",
    )
    p(
        doc,
        "11.3. ZENN no garantiza el uptime (disponibilidad continua) de servicios de terceros (Vercel, VPS, MongoDB, Firebase, Bancard, DNS, ISP, etc.). Las caídas, lentitud o incidentes de esos proveedores son ajenos a la responsabilidad de ZENN, sin perjuicio de la asistencia de soporte que el CLIENTE contrate por demanda para ayudar a diagnosticar o restablecer configuraciones a su cargo.",
    )

    heading(doc, "CLÁUSULA DÉCIMA SEGUNDA – OBLIGACIONES DEL CLIENTE")
    bullet(doc, "Proveer en tiempo y forma: logo, paleta de colores, tipografías, textos, datos de productos y demás insumos de marca.")
    bullet(doc, "Proveer dominio, correo SMTP, alta/autorización de Bancard y accesos cloud necesarios.")
    bullet(doc, "Designar un autorizador de requerimientos y de consumo de horas.")
    bullet(doc, "Realizar los pagos en los plazos convenidos.")
    bullet(doc, "Respetar la licencia de uso propio y las restricciones de no comercialización a terceros.")

    heading(doc, "CLÁUSULA DÉCIMA TERCERA – OBLIGACIONES DE ZENN")
    bullet(doc, "Ejecutar la implementación conforme a las Fases y al Anexo A.")
    bullet(doc, "Entregar el código fuente en repositorio GitHub y transferir accesos al cierre.")
    bullet(doc, "Corregir bugs en garantía conforme a la Cláusula Octava.")
    bullet(doc, "Prestar soporte y mejoras post-proyecto conforme a la tarifa y condiciones de la Cláusula Novena, cuando sean requeridos y autorizados.")
    bullet(doc, "Guardar confidencialidad sobre credenciales y datos del CLIENTE a los que acceda con motivo del Contrato.")

    heading(doc, "CLÁUSULA DÉCIMA CUARTA – TERMINACIÓN DEL VÍNCULO DE SOPORTE / MANTENIMIENTO")
    p(
        doc,
        "14.1. Cualquiera de las Partes podrá dar por terminada la relación de mantenimiento, soporte o servicios horarios posteriores, mediante preaviso escrito de treinta (30) días corridos.",
    )
    p(
        doc,
        "14.2. Durante el preaviso, el CLIENTE podrá solicitar capacitación de traspaso / handoff para que un tercero o su personal asuma el manejo del código fuente. Dicha capacitación no está incluida sin cargo y se cobrará a USD 15 por hora, previa autorización.",
    )
    p(
        doc,
        "14.3. Terminada la relación, el CLIENTE conservará el Software y el código fuente bajo la licencia de uso de este Contrato, pudiendo continuar con otro desarrollador. ZENN no estará obligado a prestar más servicios, salvo nuevo acuerdo.",
    )

    heading(doc, "CLÁUSULA DÉCIMA QUINTA – CANCELACIÓN ANTICIPADA DEL PROYECTO DE IMPLEMENTACIÓN")
    p(
        doc,
        "15.1. Si el CLIENTE cancela el proyecto antes de finalizar la Fase 2, ZENN retendrá las sumas ya abonadas en concepto de trabajo realizado, y entregará al CLIENTE el avance del Software y código fuente desarrollado hasta la fecha de cancelación, en el estado en que se encuentre.",
    )
    p(
        doc,
        "15.2. Las sumas no facturadas por hitos no alcanzados no serán exigibles, sin perjuicio de facturar trabajos extraordinarios autorizados fuera del precio cerrado, si los hubiere.",
    )

    heading(doc, "CLÁUSULA DÉCIMA SEXTA – CONFIDENCIALIDAD Y DATOS")
    p(
        doc,
        "16.1. Las Partes se obligan a mantener confidencialidad sobre información técnica, comercial, credenciales y datos personales a los que accedan con motivo de este Contrato.",
    )
    p(
        doc,
        "16.2. El CLIENTE es responsable del cumplimiento de la normativa aplicable en materia de datos personales respecto de sus usuarios finales, así como de su relación contractual con Bancard y demás proveedores.",
    )

    heading(doc, "CLÁUSULA DÉCIMA SÉPTIMA – LIMITACIÓN DE RESPONSABILIDAD")
    p(
        doc,
        "17.1. La responsabilidad agregada de ZENN derivada del presente Contrato no excederá, en ningún caso, el monto total efectivamente pagado por el CLIENTE a ZENN bajo este Contrato.",
    )
    p(
        doc,
        "17.2. ZENN no será responsable por lucro cesante, pérdida de chance, pérdida de datos imputable a terceros, daños indirectos o expectativas de ventas no alcanzadas.",
    )

    heading(doc, "CLÁUSULA DÉCIMA OCTAVA – REQUERIMIENTOS NUEVOS (HOJA DE CAMBIOS)")
    p(
        doc,
        "18.1. Cualquier funcionalidad o cambio no incluido en el Anexo A deberá cargarse en una hoja/planilla de requerimientos (o documento equivalente), estimarse en horas e implementarse: (i) con cargo a la bolsa de treinta (30) horas, mientras exista saldo y vigencia; o (ii) a USD 15 por hora; o (iii) mediante presupuesto cerrado específico aceptado por el CLIENTE.",
    )

    heading(doc, "CLÁUSULA DÉCIMA NOVENA – LEGISLACIÓN Y JURISDICCIÓN")
    p(
        doc,
        "19.1. Este Contrato se rige por las leyes de la República del Paraguay. Para cualquier controversia, las Partes se someten a los tribunales ordinarios de la ciudad de Asunción, con renuncia a cualquier otro fuero que pudiera corresponderles.",
    )

    heading(doc, "CLÁUSULA VIGÉSIMA – DISPOSICIONES FINALES")
    p(
        doc,
        "20.1. Los Anexos forman parte inseparable del Contrato. En caso de conflicto entre un Anexo y el cuerpo principal, prevalecerá el cuerpo principal, salvo que el Anexo indique expresamente lo contrario para un punto técnico de alcance.",
    )
    p(
        doc,
        "20.2. Cualquier modificación al Contrato deberá realizarse por escrito y con aceptación de ambas Partes.",
    )
    p(
        doc,
        "20.3. La nulidad de una cláusula no afectará la validez del resto del Contrato.",
    )
    p(
        doc,
        "20.4. Las Partes firman el presente Contrato en dos (2) ejemplares de un mismo tenor y a un solo efecto, en la fecha indicada en el exordio.",
        space_after=20,
    )

    # FIRMAS
    p(doc, "En prueba de conformidad, firman:", space_after=30)

    # Two column-like signature blocks using a table without borders feel
    table = doc.add_table(rows=4, cols=2)
    table.autofit = True
    c00 = table.cell(0, 0).paragraphs[0]
    c00.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c00.add_run("_______________________________")
    set_run(r)
    c01 = table.cell(0, 1).paragraphs[0]
    c01.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c01.add_run("_______________________________")
    set_run(r)

    c10 = table.cell(1, 0).paragraphs[0]
    c10.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c10.add_run("ZENN E.A.S. UNIPERSONAL")
    set_run(r, bold=True, size=10)
    c11 = table.cell(1, 1).paragraphs[0]
    c11.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c11.add_run("[RAZÓN SOCIAL DEL CLIENTE]")
    set_run(r, bold=True, size=10)

    c20 = table.cell(2, 0).paragraphs[0]
    c20.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c20.add_run("Josías Nicolás Machado Rolón")
    set_run(r, size=10)
    c21 = table.cell(2, 1).paragraphs[0]
    c21.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c21.add_run("[Nombre del Representante]")
    set_run(r, size=10)

    c30 = table.cell(3, 0).paragraphs[0]
    c30.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c30.add_run("RUC 80162448-7 — CI 5799267\nPROVEEDOR")
    set_run(r, size=9)
    c31 = table.cell(3, 1).paragraphs[0]
    c31.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c31.add_run("RUC [ ] — CI [ ]\nCLIENTE")
    set_run(r, size=9)

    # ANEXO A
    doc.add_page_break()
    p(doc, "ANEXO A", size=14, bold=True, align="center", color=PRIMARY, space_after=2)
    p(
        doc,
        "INVENTARIO DE FUNCIONALIDADES DEL SOFTWARE",
        size=12,
        bold=True,
        align="center",
        space_after=4,
    )
    p(
        doc,
        "“Plataforma E-commerce + ERP (base Zenn)”",
        size=11,
        italic=True,
        align="center",
        space_after=10,
    )
    p(
        doc,
        "El presente Anexo describe el alcance funcional incluido en el precio de USD 5.000. Todo lo no listado aquí se considera fuera de alcance y se gestiona por bolsa de horas, tarifa horaria o presupuesto específico.",
        space_after=10,
    )

    heading(doc, "A.1. Tienda online (cliente final)")
    para = p(doc, "", space_after=4)
    add_mixed(para, [("Página principal y navegación", True)])
    bullet(doc, "Home con banners, carruseles de marcas y secciones de productos destacados configurables.")
    bullet(doc, "Menú de categorías, subcategorías y sub-subcategorías (desktop y móvil).")
    bullet(doc, "Versión optimizada para celular; botón de WhatsApp; página institucional (“Nosotros”).")

    para = p(doc, "", space_after=4)
    add_mixed(para, [("Catálogo de productos", True)])
    bullet(doc, "Organización por categorías/subcategorías; ficha con galería, zoom, precio, stock, especificaciones y relacionados.")
    bullet(doc, "URLs amigables y optimización SEO básica.")

    para = p(doc, "", space_after=4)
    add_mixed(para, [("Buscador y filtros", True)])
    bullet(doc, "Buscador con vista previa; búsqueda avanzada; filtros por marca, especificaciones y precio.")

    para = p(doc, "", space_after=4)
    add_mixed(para, [("Carrito de compras", True)])
    bullet(doc, "Carrito usable sin login; contador visible; ajuste de cantidades y totales.")

    para = p(doc, "", space_after=4)
    add_mixed(para, [("Cobro / Pago – Bancard (Fase 2)", True)])
    bullet(doc, "Cobro online con tarjeta (Bancard); tarjeta nueva; tarjetas guardadas; pago con saldo/billetera.")
    bullet(doc, "Facturación, RUC e IVA; pantallas de éxito/cancelación; rollback administrativo.")

    para = p(doc, "", space_after=4)
    add_mixed(para, [("Ubicación / entrega", True)])
    bullet(doc, "Carga de dirección en perfil y checkout; selector de ubicación / geolocalización.")

    para = p(doc, "", space_after=4)
    add_mixed(para, [("Perfil del cliente", True)])
    bullet(doc, "Datos personales, foto, dirección, tarjetas, saldo, historial, favoritos, configuración.")
    bullet(doc, "Registro, login y recuperación de contraseña por email.")

    para = p(doc, "", space_after=4)
    add_mixed(para, [("Post-venta", True)])
    bullet(doc, "Detalle de pedido; seguimiento de entrega; emails automáticos; calificación del pedido.")

    heading(doc, "A.2. Panel de administración (ERP / back office)")
    bullet(doc, "Dashboard financiero, reportes de rentabilidad/márgenes, estado de cuenta, métricas anuales, análisis financiero por producto.")
    bullet(doc, "Ventas (POS interno), multi-moneda PYG/USD, IVA, RUC, adjuntos; tipos de venta, sucursales y vendedores.")
    bullet(doc, "Compras a proveedores; tipos de compra; gestión de proveedores.")
    bullet(doc, "Clientes y presupuestos (CRM) con PDF y envío por email.")
    bullet(doc, "Gestión de productos, imágenes con optimización WebP, especificaciones dinámicas, categorías, tipo de cambio USD/PYG y stock.")
    bullet(doc, "Transacciones Bancard, delivery, rollback y exportación a Excel.")
    bullet(doc, "Usuarios y permisos (ROOT/ADMIN/GENERAL) granulares por módulo.")
    bullet(doc, "Configuración centralizada de parámetros operativos.")

    heading(doc, "A.3. Automatización y carga de productos")
    bullet(doc, "Lectura/sincronización automática desde Visão Vip (categorías, specs, imágenes, precios, stock, margen).")
    bullet(doc, "Sincronización de inventario por CSV (altas, bajas, precios, visibilidad).")
    bullet(doc, "Exportación masiva a Excel (con imágenes en ZIP) y utilitarios de mantenimiento.")

    heading(doc, "A.4. Herramientas comerciales y marketing")
    bullet(doc, "Catálogo PDF; PDF de presupuestos y reportes; editor de imágenes de productos.")
    bullet(doc, "Feed de productos para Meta/Facebook; Meta Pixel; Google Analytics; emails transaccionales.")

    heading(doc, "A.5. Servicios incluidos en el precio cerrado (además del Software)")
    bullet(doc, "Implementación Fase 1 y Fase 2; diseño de marca aplicado; montaje/configuración inicial y transferencia de accesos.")
    bullet(doc, "Una (1) licencia de uso propio conforme a la Cláusula Tercera; entrega de código fuente en GitHub.")
    bullet(doc, "Diez (10) horas de capacitación; treinta (30) horas de desarrollo adicional (features fuera del inventario o diferenciaciones).")
    bullet(doc, "Cuarenta (40) días de garantía de bugs post-aceptación de Fase 2.")

    heading(doc, "A.6. Fuera de alcance (se cotiza / horas)")
    bullet(doc, "Cualquier módulo o función no listada en este Anexo.")
    bullet(doc, "App móvil nativa; desarrollos a medida no previstos; integraciones nuevas no contempladas.")
    bullet(doc, "Campañas publicitarias, community management, producción de contenido, y costos de infraestructura de terceros.")
    bullet(doc, "Soporte y mantenimiento posteriores a la garantía, a USD 15/hora por demanda.")

    p(doc, "", space_after=16)
    p(
        doc,
        "Las Partes dejan constancia de haber leído y aceptado el presente Anexo A en la misma fecha del Contrato.",
        space_after=24,
    )

    table = doc.add_table(rows=3, cols=2)
    for col, title in enumerate(["ZENN E.A.S. UNIPERSONAL", "[CLIENTE – RAZÓN SOCIAL]"]):
        para = table.cell(0, col).paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run("_______________________________")
        set_run(r)
        para2 = table.cell(1, col).paragraphs[0]
        para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para2.add_run(title)
        set_run(r, bold=True, size=10)
        para3 = table.cell(2, col).paragraphs[0]
        para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para3.add_run("Firma y aclaración")
        set_run(r, size=9, italic=True)

    # ANEXO B - modelo acta
    doc.add_page_break()
    p(doc, "ANEXO B", size=14, bold=True, align="center", color=PRIMARY, space_after=2)
    p(doc, "MODELO DE ACTA DE ACEPTACIÓN DE FASE", size=12, bold=True, align="center", space_after=12)

    p(
        doc,
        "Acta de Aceptación – Fase N.º ____ (1 / 2)",
        bold=True,
        space_after=10,
    )
    p(
        doc,
        "Fecha: ____ / ____ / ________    Proyecto: Plataforma E-commerce + ERP (base Zenn)",
        space_after=8,
    )
    p(
        doc,
        "El CLIENTE declara haber revisado los entregables de la Fase indicada, conforme al Contrato y al Anexo A, y manifiesta (marcar una opción):",
        space_after=6,
    )
    bullet(doc, "ACEPTACIÓN TOTAL: la Fase se considera concluida a satisfacción.")
    bullet(doc, "ACEPTACIÓN CON OBSERVACIONES: se detallan bugs/errores de alcance a corregir (no incluye nuevas funciones):")
    p(doc, "1. _______________________________________________________________", space_after=4)
    p(doc, "2. _______________________________________________________________", space_after=4)
    p(doc, "3. _______________________________________________________________", space_after=10)
    p(
        doc,
        "Plazo de revisión según Contrato: quince (15) días corridos desde la notificación de ZENN. Sin respuesta, aceptación tácita.",
        space_after=20,
    )

    table = doc.add_table(rows=3, cols=2)
    for col, title in enumerate(["Por ZENN", "Por el CLIENTE"]):
        para = table.cell(0, col).paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run("_______________________________")
        set_run(r)
        para2 = table.cell(1, col).paragraphs[0]
        para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para2.add_run(title)
        set_run(r, bold=True, size=10)
        para3 = table.cell(2, col).paragraphs[0]
        para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para3.add_run("Nombre, CI y firma")
        set_run(r, size=9, italic=True)

    doc.save(OUT)
    print("OK", OUT)


if __name__ == "__main__":
    main()
