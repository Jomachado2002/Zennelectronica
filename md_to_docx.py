#!/usr/bin/env python3
"""Convierte la propuesta comercial (Markdown) en un .docx bien formateado."""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "/Users/josiasnicolas02gmail.com/Documents/Zennelectronica/PROPUESTA_COMERCIAL_SOFTWARE.md"
OUT = "/Users/josiasnicolas02gmail.com/Documents/Zennelectronica/PROPUESTA_COMERCIAL_SOFTWARE.docx"

PRIMARY = RGBColor(0x0B, 0x3D, 0x91)   # azul corporativo
DARK = RGBColor(0x1F, 0x2A, 0x37)      # gris oscuro texto
ACCENT = RGBColor(0x0E, 0x7C, 0x3A)    # verde acento
HEADER_BG = "0B3D91"
ZEBRA_BG = "EEF3FB"

BASE_FONT = "Calibri"


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_inline(paragraph, text, base_color=DARK, base_size=11, base_bold=False):
    """Agrega runs manejando **negrita** y `codigo`."""
    tokens = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for tok in tokens:
        if not tok:
            continue
        bold = base_bold
        mono = False
        content = tok
        if tok.startswith('**') and tok.endswith('**'):
            bold = True
            content = tok[2:-2]
        elif tok.startswith('`') and tok.endswith('`'):
            mono = True
            content = tok[1:-1]
        run = paragraph.add_run(content)
        run.bold = bold
        run.font.color.rgb = base_color
        run.font.size = Pt(base_size)
        run.font.name = 'Consolas' if mono else BASE_FONT
    return paragraph


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    if level == 1:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = PRIMARY
        run.font.name = BASE_FONT
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(15)
        run.font.color.rgb = PRIMARY
        run.font.name = BASE_FONT
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        # borde inferior
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '2')
        bottom.set(qn('w:color'), 'B9CCE8')
        pbdr.append(bottom)
        pPr.append(pbdr)
    else:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12.5)
        run.font.color.rgb = ACCENT
        run.font.name = BASE_FONT
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    return p


def build_table(doc, rows):
    header = rows[0]
    body = rows[2:]  # rows[1] es separador ---
    ncol = len(header)
    table = doc.add_table(rows=1, cols=ncol)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        set_cell_bg(hdr_cells[i], HEADER_BG)
        para = hdr_cells[i].paragraphs[0]
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.space_before = Pt(2)
        run = para.add_run(h.replace('**', '').strip())
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10.5)
        run.font.name = BASE_FONT
    for r, row in enumerate(body):
        cells = table.add_row().cells
        for i in range(ncol):
            val = row[i] if i < len(row) else ''
            if r % 2 == 1:
                set_cell_bg(cells[i], ZEBRA_BG)
            para = cells[i].paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            add_inline(para, val.strip(), base_color=DARK, base_size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def parse_table_row(line):
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [c.strip() for c in line.split('|')]


def main():
    with open(SRC, encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = BASE_FONT
    style.font.size = Pt(11)
    style.font.color.rgb = DARK
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Regla horizontal
        if re.match(r'^-{3,}$', stripped):
            i += 1
            continue

        # Encabezados
        m = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if m:
            level = len(m.group(1))
            add_heading(doc, m.group(2).strip(), level)
            i += 1
            continue

        # Tabla
        if stripped.startswith('|'):
            tbl_lines = []
            while i < n and lines[i].strip().startswith('|'):
                tbl_lines.append(parse_table_row(lines[i]))
                i += 1
            if len(tbl_lines) >= 2:
                build_table(doc, tbl_lines)
            continue

        # Blockquote
        if stripped.startswith('>'):
            content = stripped.lstrip('>').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '18')
            left.set(qn('w:space'), '8')
            left.set(qn('w:color'), '0E7C3A')
            pbdr.append(left)
            pPr.append(pbdr)
            add_inline(p, content, base_color=RGBColor(0x37, 0x47, 0x51), base_size=10.5)
            for r in p.runs:
                r.italic = True
            i += 1
            continue

        # Lista numerada
        mnum = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if mnum:
            p = doc.add_paragraph(style='List Number')
            add_inline(p, mnum.group(2).strip())
            i += 1
            continue

        # Lista con viñetas
        mbul = re.match(r'^[-*]\s+(.*)$', stripped)
        if mbul:
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, mbul.group(1).strip())
            i += 1
            continue

        # Párrafo normal
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        add_inline(p, stripped)
        i += 1

    doc.save(OUT)
    print("OK ->", OUT)


if __name__ == '__main__':
    main()
